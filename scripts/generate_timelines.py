from pathlib import Path
import re
from docx import Document

# --- Paths ---
SRC = Path("/Users/adrienallorant/Documents/Jeanne_Jozon")
SITE = Path("/Users/adrienallorant/Documents/hugo_site_jj/content")

# --- Load text helper ---
def extract_year_blocks(doc_text):
    """Return list of (year, textblock) tuples."""
    pattern = r"(?P<year>18\d{2}|19\d{2}|20\d{2})(?:[\s\n\-:]+)(?P<text>.*?)(?=(?:18\d{2}|19\d{2}|20\d{2}|$))"
    matches = re.finditer(pattern, doc_text, flags=re.S)
    return [(m.group("year"), re.sub(r"\s+", " ", m.group("text")).strip()) for m in matches]

bio_doc = Document(SRC / "Biographie" / "Biographie.docx")
salons_doc = Document(SRC / "Salons" / "Participation de Jeanne Jozon aux expositions artistiques.docx")

# --- Biographie ---
bio_text = "\n".join([p.text for p in bio_doc.paragraphs])
bio_blocks = extract_year_blocks(bio_text)

bio_md = [
    "---",
    'title: "Biographie"',
    "description: Parcours de Jeanne Jozon (1868–1946)",
    "---",
    "",
    "# Biographie de Jeanne Jozon",
    "",
    '<div class="life-timeline">',
]

for year, text in bio_blocks:
    bio_md.append(f'<div class="event"><div class="date">{year}</div><p>{text}</p></div>')

bio_md.append("</div>\n")

# Add style block
bio_md.append("""
<style>
.life-timeline{display:flex;overflow-x:auto;scroll-snap-type:x mandatory;gap:2rem;padding:2rem 0;background:linear-gradient(to right,#f7f3eb,#fff);}
.event{position:relative;flex:0 0 240px;scroll-snap-align:start;background:#fff;border:2px solid #c07c3a;border-radius:20px;padding:1rem;text-align:center;box-shadow:0 4px 10px rgba(0,0,0,0.1);}
.event .date{font-family:"Cormorant Garamond",serif;font-weight:600;color:#7b563a;margin-bottom:0.3rem;}
.event p{font-size:0.9rem;line-height:1.3;}
</style>
""")

# Add full biographical text below
bio_md.append("\n---\n\n## Texte biographique\n\n")
bio_md.append(bio_text)

(SITE / "biographie" / "_index.md").write_text("\n".join(bio_md), encoding="utf-8")
print("✓ Biographie timeline generated.")

# --- Salons ---
salons_text = "\n".join([p.text for p in salons_doc.paragraphs])
salons_blocks = extract_year_blocks(salons_text)

salons_md = [
    "---",
    'title: "Salons et expositions"',
    "description: Participation de Jeanne Jozon aux expositions artistiques (1895–1939)",
    "---",
    "",
    "# Participation aux salons et expositions",
    "",
    '<div class="life-timeline">',
]

for year, text in salons_blocks:
    salons_md.append(f'<div class="event"><div class="date">{year}</div><p>{text}</p></div>')

salons_md.append("</div>\n")
salons_md.append("""
<style>
.life-timeline{display:flex;overflow-x:auto;scroll-snap-type:x mandatory;gap:2rem;padding:2rem 0;background:linear-gradient(to right,#f7f3eb,#fff);}
.event{position:relative;flex:0 0 260px;scroll-snap-align:start;background:#fff;border:2px solid #c07c3a;border-radius:20px;padding:1rem;text-align:center;box-shadow:0 4px 10px rgba(0,0,0,0.1);}
.event .date{font-family:"Cormorant Garamond",serif;font-weight:600;color:#7b563a;margin-bottom:0.3rem;}
.event p{font-size:0.9rem;line-height:1.3;}
</style>
""")

salons_md.append("\n---\n\n## Texte explicatif\n\n")
salons_md.append(salons_text)

(SITE / "salons" / "_index.md").write_text("\n".join(salons_md), encoding="utf-8")
print("✓ Salons timeline generated.")